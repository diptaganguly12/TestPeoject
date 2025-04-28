from flask import Flask, request, jsonify, render_template, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename
from models import DataStore, CCTVCamera, Schedule
from cctv_manager import CCTVManager
from schedule_manager import ScheduleManager
import os
import datetime
import uuid
import json
import ollama
import atexit
import pytz
import time
from hospital_data import HospitalDataManager
import pandas as pd
from analysis_queue import AnalysisQueue
from werkzeug.middleware.proxy_fix import ProxyFix
import logging
from logging.handlers import RotatingFileHandler
from location_validator import is_location_valid
from hospital_locations import HospitalLocations
from functools import wraps
from flask import session, redirect, url_for
from user_manager import UserManager
from flask_jwt_extended import JWTManager, create_access_token, get_jwt_identity, jwt_required, verify_jwt_in_request

app = Flask(__name__)
CORS(app)

# Setup logging
if not os.path.exists('logs'):
    os.makedirs('logs')

handler = RotatingFileHandler('logs/app.log', maxBytes=10000000, backupCount=5)
handler.setLevel(logging.INFO)
app.logger.addHandler(handler)
app.logger.setLevel(logging.INFO)

# Configure for proxy
app.wsgi_app = ProxyFix(
    app.wsgi_app,
    x_for=1,
    x_proto=1,
    x_host=1,
    x_prefix=1
)

# Configuration
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
DATA_FOLDER = os.path.join(BASE_DIR, 'data')
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'PNG'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size
hospital_locations = HospitalLocations(DATA_FOLDER)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Add secret key for session management
app.secret_key = 'your-secure-secret-key-here'  # Change this in production
user_manager = UserManager(DATA_FOLDER)

# Configure JWT
app.config['JWT_SECRET_KEY'] = 'your-jwt-secret-key-change-in-production'  # Change this in production
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = datetime.timedelta(days=1)
jwt = JWTManager(app)

# Create required directories
for folder in [UPLOAD_FOLDER, DATA_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

def get_current_time():
    """Get current time in UTC format YYYY-MM-DD HH:MM:SS"""
    return datetime.datetime.now(pytz.UTC).strftime("%Y-%m-%d %H:%M:%S")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper to get username (either from session or JWT)
def get_current_username():
    if 'username' in session:
        return session['username']
    try:
        return get_jwt_identity()
    except Exception:
        return None

# Authentication decorators
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# New dual authentication decorator that accepts both session and token
def dual_auth_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Try to verify JWT token first
        try:
            verify_jwt_in_request()
            return f(*args, **kwargs)
        except Exception:
            # If JWT verification fails, check for session authentication
            if 'username' not in session:
                if request.content_type == 'application/json' or request.headers.get('Accept') == 'application/json':
                    return jsonify({'success': False, 'error': 'Unauthorized access'}), 401
                return redirect(url_for('login'))
            return f(*args, **kwargs)
    return decorated_function

@app.before_request
def log_request():
    app.logger.info(f'Request: {request.method} {request.url} from {request.remote_addr}')
    app.logger.info(f'Headers: {dict(request.headers)}')

@app.after_request
def log_response(response):
    app.logger.info(f'Response: {response.status_code}')
    return response

# Authentication routes
@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'username' in session:
        return redirect(url_for('index'))
    
    error = None
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        
        if user_manager.validate_user(username, password):
            session['username'] = username
            return redirect(url_for('index'))
        else:
            error = 'Invalid username or password'
    
    return render_template('login.html', error=error)

# New API for token-based login
@app.route('/api/login', methods=['POST'])
def api_login():
    if not request.is_json:
        return jsonify({"success": False, "error": "Missing JSON in request"}), 400
    
    username = request.json.get('username', None)
    password = request.json.get('password', None)
    
    if not username or not password:
        return jsonify({"success": False, "error": "Missing username or password"}), 400
    
    if user_manager.validate_user(username, password):
        # Create access token
        access_token = create_access_token(identity=username)
        return jsonify({
            "success": True,
            "access_token": access_token,
            "username": username
        })
    else:
        return jsonify({"success": False, "error": "Invalid username or password"}), 401

@app.route('/register', methods=['GET', 'POST'])
def register():
    if 'username' in session:
        return redirect(url_for('index'))
    
    error = None
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')
        
        if password != confirm_password:
            error = 'Passwords do not match'
        else:
            success, message = user_manager.register_user(username, password)
            if success:
                return redirect(url_for('login'))
            else:
                error = message
    
    return render_template('register.html', error=error)

# API for mobile app registration
@app.route('/api/register', methods=['POST'])
def api_register():
    if not request.is_json:
        return jsonify({"success": False, "error": "Missing JSON in request"}), 400
    
    username = request.json.get('username', None)
    password = request.json.get('password', None)
    confirm_password = request.json.get('confirm_password', None)
    
    if not username or not password or not confirm_password:
        return jsonify({"success": False, "error": "Missing required fields"}), 400
    
    if password != confirm_password:
        return jsonify({"success": False, "error": "Passwords do not match"}), 400
    
    success, message = user_manager.register_user(username, password)
    if success:
        return jsonify({"success": True, "message": "Registration successful"})
    else:
        return jsonify({"success": False, "error": message}), 400

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

# Main application routes
@app.route('/')
@login_required
def index():
    floors = sorted(list(set(camera.floor for camera in data_store.cameras.values())))
    return render_template('index.html',
                         current_time=get_current_time(),
                         user_login=session['username'],
                         floors=floors)

@app.route('/api/system-info')
@dual_auth_required
def get_system_info():
    current_user = get_current_username()
    return jsonify({
        'success': True,
        'current_time': get_current_time(),
        'user_login': current_user
    })

# Hospital data routes
@app.route('/api/hospitals', methods=['GET'])
@dual_auth_required
def get_hospitals():
    try:
        hospitals = hospital_manager.get_hospitals()
        return jsonify({'success': True, 'data': hospitals})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/buildings/<hospital>', methods=['GET'])
@dual_auth_required
def get_buildings(hospital):
    try:
        buildings = hospital_manager.get_buildings(hospital)
        return jsonify({'success': True, 'data': buildings})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/floors/<hospital>/<building>', methods=['GET'])
@dual_auth_required
def get_floors(hospital, building):
    try:
        floors = hospital_manager.get_floors(hospital, building)
        return jsonify({'success': True, 'data': floors})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/categories/<hospital>/<building>/<floor>', methods=['GET'])
@dual_auth_required
def get_categories(hospital, building, floor):
    try:
        categories = hospital_manager.get_categories(hospital, building, floor)
        return jsonify({'success': True, 'data': categories})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/uploads/<filename>')
@dual_auth_required
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

# Image analysis related functions
def analyze_image(image_path):
    """Use Ollama vision model for image analysis with retries and timeout"""
    try:
        prompt = """You are an AI detector. Analyze the cleanliness of the given image based on 3 parameters: cleanliness, hygiene, and pest control.
        Note: This analysis is for government hospitals. Be understanding of resource constraints and only give lower scores if there are major stains, dirt, or clear problems visible.

        Return ONLY a JSON object with the following structure:
        {
          "category": "WashRoom/Floor/Ward/emergency/lab/ICU/OT.... You decide",
          "timestamp": "YYYY-MM-DD HH:MM:SS",
          "analysis_results": {
            "cleanliness": {
              "score": ?%,
              "status": "IF the score is more than 80 then is Good,between 70 and 80 then Acceptable,less than 70 Needs Attention"
            },
            "hygiene": {
              "score": ?%,
              "status": "IF the score is more than 80 then is Good,between 70 and 80 then Acceptable,less than 70 Needs Attention"
            },
            "pest_control": {
              "score": ?%,
              "status": "IF the score is more than 80 then is Good,between 70 and 80 then Acceptable,less than 70 Needs Attention"
            },
            "overall": {
              "score": ?%,
              "status": "IF the score is more than 80 then is Good,between 70 and 80 then Acceptable,less than 70 Needs Attention"
            }
          },
          "comment": "Brief comment on the image explain the analysis in short."
        }"""

        max_retries = 3
        retry_delay = 2

        for attempt in range(max_retries):
            try:
                response = ollama.chat(
                    model="llama3.2-vision:11b",
                    messages=[
                        {"role": "system", "content": "You are an AI detector that strictly returns JSON formatted analysis."},
                        {"role": "user", "content": prompt, "images": [image_path]},
                    ]
                )

                response_text = response.get("message", {}).get("content", "{}").strip()
                
                try:
                    parsed_response = json.loads(response_text)
                    result = {
                        'cleanliness': parsed_response['analysis_results']['cleanliness'],
                        'hygiene': parsed_response['analysis_results']['hygiene'],
                        'pest_control': parsed_response['analysis_results']['pest_control'],
                        'overall': parsed_response['analysis_results']['overall'],
                        'comment': parsed_response.get('comment', 'No comment available')
                    }
                    return result

                except json.JSONDecodeError:
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                        continue
                    print(f"Failed to parse JSON response after {max_retries} attempts")
                    return None

            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(retry_delay)
                    continue
                print(f"Error during Ollama analysis after {max_retries} attempts: {str(e)}")
                return None

    except Exception as e:
        print(f"Critical error in analyze_image: {str(e)}")
        return None

# Initialize managers
data_store = DataStore(data_dir=DATA_FOLDER)
cctv_manager = CCTVManager(data_store, upload_dir=UPLOAD_FOLDER)
schedule_manager = ScheduleManager(data_store, cctv_manager, analyze_image)
hospital_manager = HospitalDataManager(DATA_FOLDER)
analysis_queue = AnalysisQueue(UPLOAD_FOLDER, analyze_image)

# Start services
schedule_manager.start()
analysis_queue.start()

# Image upload and analysis routes
@app.route('/api/upload', methods=['POST'])
@dual_auth_required
def upload_image():
    if 'image' not in request.files:
        return jsonify({'success': False, 'error': 'No image file provided'}), 400
    
    file = request.files['image']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No selected file'}), 400

    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file type'}), 400

    try:
        geolocation = request.form.get('geolocation')
        hospital_name = request.form.get('hospital_name')
        
        if not geolocation:
            return jsonify({'success': False, 'error': 'Geolocation data is required'}), 400
        
        if not hospital_name:
            return jsonify({'success': False, 'error': 'Hospital name is required'}), 400
        
        try:
            geo_data = json.loads(geolocation)
            user_lat = geo_data.get('latitude')
            user_lon = geo_data.get('longitude')
            
            if not user_lat or not user_lon:
                return jsonify({'success': False, 'error': 'Invalid geolocation data'}), 400
            
            is_valid, error_message = is_location_valid(
                user_lat, 
                user_lon, 
                hospital_name,
                hospital_locations.data
            )
            
            if not is_valid:
                return jsonify({'success': False, 'error': error_message}), 403
                
        except json.JSONDecodeError:
            return jsonify({'success': False, 'error': 'Invalid geolocation format'}), 400

        timestamp = get_current_time()
        safe_timestamp = timestamp.replace(':', '').replace(' ', '_')
        filename = secure_filename(f"{safe_timestamp}_{uuid.uuid4().hex[:8]}_{file.filename}")
        filename = filename.replace(' ', '_').replace('(', '').replace(')', '')
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
        file.save(filepath)

        if not os.path.exists(filepath):
            raise Exception(f"Failed to save file at {filepath}")

        building = request.form.get('building', 'Unknown')
        floor = request.form.get('floor', 'Unknown')
        category = request.form.get('category', 'Unknown')
        location = request.form.get('location', 'Unknown')
        
        # Get priority flag from request, default to False
        priority = request.form.get('priority', 'false').lower() == 'true'
        
        # New flag to process immediately
        run_immediately = request.form.get('run_immediately', 'false').lower() == 'true'
        
        # If priority is true and run_immediately is true, process the image right away
        if priority and run_immediately:
            # Analyze the image immediately
            analysis_results = analyze_image(filepath)
            
            if analysis_results:
                metadata = {
                    'filename': filename,
                    'hospital_name': hospital_name,
                    'building': building,
                    'floor': floor,
                    'category': category,
                    'location': location,
                    'timestamp': timestamp,
                    'geolocation': geolocation,
                    'status': 'completed',
                    'priority': priority,
                    'analysis_results': analysis_results,
                    'comment': analysis_results.get('comment', 'No comment available'),
                    'task_id': None,
                    'created_at': timestamp,
                    'completed_at': timestamp,
                    'uploaded_by': get_current_username()
                }
                
                metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename}.json")
                with open(metadata_path, 'w') as f:
                    json.dump(metadata, f, indent=2)
                
                return jsonify({
                    'success': True,
                    'data': metadata,
                    'message': 'Image uploaded and processed immediately'
                })
            else:
                # If analysis failed, still queue it
                app.logger.warning(f"Immediate analysis failed for {filename}, falling back to queue")
        
        # Log priority parameter for debugging
        app.logger.info(f"Priority parameter: {request.form.get('priority')} (interpreted as {priority})")
        app.logger.info(f"Run immediately parameter: {request.form.get('run_immediately')} (interpreted as {run_immediately})")

        metadata = {
            'filename': filename,
            'hospital_name': hospital_name,
            'building': building,
            'floor': floor,
            'category': category,
            'location': location,
            'timestamp': timestamp,
            'geolocation': geolocation,
            'status': 'pending',
            'priority': priority,
            'analysis_results': None,
            'comment': None,
            'task_id': None,
            'created_at': timestamp,
            'completed_at': None,
            'uploaded_by': get_current_username()
        }

        metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{filename}.json")
        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)

        # Pass priority flag to add_task
        task_id = analysis_queue.add_task(filename, metadata, priority)
        metadata['task_id'] = task_id

        with open(metadata_path, 'w') as f:
            json.dump(metadata, f, indent=2)

        # Include priority information in response
        return jsonify({
            'success': True,
            'data': metadata,
            'message': 'Image upload successful' + (' (priority processing)' if priority else ' (scheduled for processing after 9 PM)')
        })

    except Exception as e:
        print(f"Upload error: {str(e)}")
        if 'filepath' in locals() and os.path.exists(filepath):
            try:
                os.remove(filepath)
            except:
                pass
        if 'metadata_path' in locals() and os.path.exists(metadata_path):
            try:
                os.remove(metadata_path)
            except:
                pass
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/task/<task_id>/priority', methods=['POST'])
@dual_auth_required
def update_task_priority(task_id):
    """Update the priority of an existing task"""
    if not request.is_json:
        return jsonify({"success": False, "error": "Missing JSON in request"}), 400
    
    priority = request.json.get('priority', True)
    run_immediately = request.json.get('run_immediately', False)  # New parameter
    
    if not isinstance(priority, bool):
        return jsonify({"success": False, "error": "Priority must be a boolean value"}), 400
    
    app.logger.info(f"Updating task {task_id} priority to {priority}, run immediately: {run_immediately}")
        
    success = analysis_queue.set_task_priority(task_id, priority)
    
    if success:
        # If successful, update the metadata file too
        updated = False
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            if filename.endswith('.json'):
                try:
                    metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    with open(metadata_path, 'r') as f:
                        metadata = json.load(f)
                        if metadata.get('task_id') == task_id:
                            metadata['priority'] = priority
                            with open(metadata_path, 'w') as f:
                                json.dump(metadata, f, indent=2)
                            updated = True
                            task_filename = metadata.get('filename')
                            break
                except Exception as e:
                    print(f"Error updating metadata file for task {task_id}: {str(e)}")
        
        # Process immediately if requested and it's a high priority task
        if run_immediately and priority:
            task = analysis_queue.get_task(task_id)
            if task and task.status == 'pending':
                # Use the existing process_task_now logic
                import threading
                
                def process_task():
                    # Set the task to processing status
                    task.status = 'processing'
                    analysis_queue._update_task_status(task.id, 'processing')
                    
                    try:
                        # Process the task
                        filepath = os.path.join(app.config['UPLOAD_FOLDER'], task.filename)
                        results = analyze_image(filepath)
                        
                        if results:
                            # Update metadata with analysis results
                            metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                            
                            try:
                                with open(metadata_path, 'r') as f:
                                    metadata = json.load(f)
                            except:
                                metadata = {}
                            
                            metadata['analysis_results'] = results
                            metadata['comment'] = results.get('comment', 'No comment available')
                            metadata['status'] = 'completed'
                            metadata['completed_at'] = get_current_time()
                            
                            with open(metadata_path, 'w') as f:
                                json.dump(metadata, f, indent=2)
                            
                            # Update task status
                            task.completed_at = get_current_time()
                            task.status = 'completed'
                            analysis_queue._update_task_status(task.id, 'completed')
                        else:
                            # Failed to analyze
                            task.status = 'failed'
                            task.error = 'Analysis failed to produce results'
                            analysis_queue._update_task_status(task.id, 'failed', 'Analysis failed to produce results')
                            
                            metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                            try:
                                with open(metadata_path, 'r') as f:
                                    metadata = json.load(f)
                                
                                metadata['status'] = 'failed'
                                metadata['error'] = 'Analysis failed to produce results'
                                
                                with open(metadata_path, 'w') as f:
                                    json.dump(metadata, f, indent=2)
                            except:
                                pass
                    except Exception as e:
                        # Handle errors
                        task.status = 'failed'
                        task.error = str(e)
                        analysis_queue._update_task_status(task.id, 'failed', str(e))
                        
                        # Update metadata
                        metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                        try:
                            with open(metadata_path, 'r') as f:
                                metadata = json.load(f)
                            
                            metadata['status'] = 'failed'
                            metadata['error'] = str(e)
                            
                            with open(metadata_path, 'w') as f:
                                json.dump(metadata, f, indent=2)
                        except:
                            pass
                
                # Start processing in a separate thread
                processing_thread = threading.Thread(target=process_task)
                processing_thread.daemon = True
                processing_thread.start()
                
                return jsonify({
                    "success": True,
                    "message": f"Task priority updated to high and processing started immediately",
                    "metadata_updated": updated,
                    "task_status": "processing"
                })
        
        return jsonify({
            "success": True,
            "message": f"Task priority updated to {'high' if priority else 'normal'}",
            "metadata_updated": updated
        })
    else:
        return jsonify({
            "success": False,
            "error": "Failed to update task priority. Task may not exist or is already processing."
        }), 404

@app.route('/api/analysis-status/<task_id>')
@dual_auth_required
def get_analysis_status(task_id):
    task = analysis_queue.get_task(task_id)
    if not task:
        return jsonify({'success': False, 'error': 'Task not found'}), 404
    
    return jsonify({
        'success': True,
        'data': {
            'task_id': task.id,
            'status': task.status,
            'priority': task.priority,  # Include priority in response
            'created_at': task.created_at,
            'completed_at': task.completed_at,
            'error': task.error,
            'filename': task.filename
        }
    })

@app.route('/api/images', methods=['GET'])
@dual_auth_required
def get_images():
    try:
        images = []
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            if filename.endswith('.json'):
                try:
                    with open(os.path.join(app.config['UPLOAD_FOLDER'], filename), 'r') as f:
                        metadata = json.load(f)
                        image_filename = metadata.get('filename')
                        if image_filename and os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], image_filename)):
                            images.append(metadata)
                except Exception as e:
                    print(f"Error reading metadata file {filename}: {str(e)}")
                    continue

        images.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
        return jsonify({'success': True, 'data': images})

    except Exception as e:
        print(f"Error getting images: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/cameras/<camera_id>/capture', methods=['POST'])
@dual_auth_required
def capture_camera(camera_id):
    if camera_id not in data_store.cameras:
        return jsonify({'success': False, 'error': 'Camera not found'}), 404

    # Get priority flag from request, default to False
    priority = request.json.get('priority', False) if request.is_json else False

    camera = data_store.cameras[camera_id]
    filepath = cctv_manager.capture_from_camera(camera)
    
    if filepath:
        # Only analyze immediately if priority is True, otherwise queue it
        if priority:
            analysis_results = analyze_image(filepath)
            if analysis_results:
                metadata = {
                    'filename': os.path.basename(filepath),
                    'location': camera.location,
                    'category': camera.floor,
                    'timestamp': get_current_time(),
                    'camera_id': camera_id,
                    'priority': priority,
                    'analysis_results': analysis_results,
                    'comment': analysis_results.get('comment', 'No comment available'),
                    'uploaded_by': get_current_username()
                }
                
                metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.basename(filepath)}.json")
                with open(metadata_path, 'w') as f:
                    json.dump(metadata, f, indent=2)
                    
                return jsonify({
                    'success': True,
                    'data': metadata
                })
        else:
            # Queue for later processing
            metadata = {
                'filename': os.path.basename(filepath),
                'location': camera.location,
                'category': camera.floor,
                'timestamp': get_current_time(),
                'camera_id': camera_id,
                'priority': priority,
                'status': 'pending',
                'analysis_results': None,
                'comment': None,
                'uploaded_by': get_current_username()
            }
            
            metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{os.path.basename(filepath)}.json")
            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
                
            # Add to queue for processing after 9PM
            task_id = analysis_queue.add_task(os.path.basename(filepath), metadata, priority)
            metadata['task_id'] = task_id
            
            with open(metadata_path, 'w') as f:
                json.dump(metadata, f, indent=2)
                
            return jsonify({
                'success': True,
                'data': metadata,
                'message': 'Image captured and queued for processing after 9 PM'
            })
    
    return jsonify({'success': False, 'error': 'Failed to capture image'}), 500

# New endpoint to check queue status and scheduler state
@app.route('/api/queue-status')
@dual_auth_required
def get_queue_status():
    try:
        # Get counts of priority and regular tasks
        priority_count = analysis_queue.priority_queue.qsize()
        regular_count = analysis_queue.queue.qsize()
        
        # Get processing state
        processing_enabled = analysis_queue.processing_enabled
        
        # Get counts by different statuses
        pending_count = 0
        processing_count = 0
        completed_count = 0
        failed_count = 0
        
        for task in analysis_queue.tasks.values():
            if task.status == 'pending':
                pending_count += 1
            elif task.status == 'processing':
                processing_count += 1
            elif task.status == 'completed':
                completed_count += 1
            elif task.status == 'failed':
                failed_count += 1
        
        # Get current time for reference
        current_time = get_current_time()
        
        # Calculate when regular processing will begin
        next_processing_time = "21:00:00" if not processing_enabled else "In progress"
        
        return jsonify({
            'success': True,
            'data': {
                'queue_counts': {
                    'priority': priority_count,
                    'regular': regular_count,
                    'total': priority_count + regular_count
                },
                'task_status': {
                    'pending': pending_count,
                    'processing': processing_count,
                    'completed': completed_count,
                    'failed': failed_count,
                    'total': len(analysis_queue.tasks)
                },
                'processing_status': {
                    'enabled': processing_enabled,
                    'next_processing_time': next_processing_time,
                    'current_time': current_time
                }
            }
        })
    except Exception as e:
        print(f"Error getting queue status: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# New endpoint to manually enable processing (for admin control)
@app.route('/api/enable-processing', methods=['POST'])
@dual_auth_required
def enable_processing():
    try:
        analysis_queue._enable_processing()
        return jsonify({
            'success': True,
            'message': 'Processing enabled manually'
        })
    except Exception as e:
        print(f"Error enabling processing: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# New endpoint to verify token and get user data
@app.route('/api/user', methods=['GET'])
@jwt_required()
def get_user_data():
    current_user = get_jwt_identity()
    return jsonify({
        'success': True,
        'username': current_user,
        'timestamp': get_current_time()
    })
    
    
@app.route('/api/hospital-location/<hospital_name>')
@dual_auth_required
def get_hospital_location(hospital_name):
    """API endpoint to get hospital location data for geo-fence verification"""
    try:
        # Get location from hospital_locations class
        location = hospital_locations.get_hospital_location(hospital_name)
        
        if location:
            return jsonify({
                'success': True,
                'data': location
            })
        else:
            # If hospital not found in database, return 404
            app.logger.error(f"Hospital location not found for: {hospital_name}")
            return jsonify({
                'success': False,
                'error': f"Hospital location data not found for {hospital_name}"
            }), 404
    except Exception as e:
        app.logger.error(f"Error getting hospital location: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500    

@app.route('/api/process-now/<task_id>', methods=['POST'])
@dual_auth_required
def process_task_now(task_id):
    """Process a specific task immediately"""
    try:
        # Get the task
        task = analysis_queue.get_task(task_id)
        if not task:
            return jsonify({'success': False, 'error': 'Task not found'}), 404
        
        # Check if task is already processing or completed
        if task.status != 'pending':
            return jsonify({
                'success': False, 
                'error': f'Task cannot be processed now. Current status: {task.status}'
            }), 400
        
        # Process the task immediately in a separate thread
        import threading
        
        def process_task():
            # Set the task to processing status
            task.status = 'processing'
            analysis_queue._update_task_status(task.id, 'processing')
            
            try:
                # Process the task
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], task.filename)
                results = analyze_image(filepath)
                
                if results:
                    # Update metadata with analysis results
                    metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                    
                    try:
                        with open(metadata_path, 'r') as f:
                            metadata = json.load(f)
                    except:
                        metadata = {}
                    
                    metadata['analysis_results'] = results
                    metadata['comment'] = results.get('comment', 'No comment available')
                    metadata['status'] = 'completed'
                    metadata['completed_at'] = get_current_time()
                    
                    with open(metadata_path, 'w') as f:
                        json.dump(metadata, f, indent=2)
                    
                    # Update task status
                    task.completed_at = get_current_time()
                    task.status = 'completed'
                    analysis_queue._update_task_status(task.id, 'completed')
                else:
                    # Failed to analyze
                    task.status = 'failed'
                    task.error = 'Analysis failed to produce results'
                    analysis_queue._update_task_status(task.id, 'failed', 'Analysis failed to produce results')
                    
                    metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                    try:
                        with open(metadata_path, 'r') as f:
                            metadata = json.load(f)
                        
                        metadata['status'] = 'failed'
                        metadata['error'] = 'Analysis failed to produce results'
                        
                        with open(metadata_path, 'w') as f:
                            json.dump(metadata, f, indent=2)
                    except:
                        pass
            except Exception as e:
                # Handle errors
                task.status = 'failed'
                task.error = str(e)
                analysis_queue._update_task_status(task.id, 'failed', str(e))
                
                # Update metadata
                metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                try:
                    with open(metadata_path, 'r') as f:
                        metadata = json.load(f)
                    
                    metadata['status'] = 'failed'
                    metadata['error'] = str(e)
                    
                    with open(metadata_path, 'w') as f:
                        json.dump(metadata, f, indent=2)
                except:
                    pass
        
        # Start processing in a separate thread
        processing_thread = threading.Thread(target=process_task)
        processing_thread.daemon = True
        processing_thread.start()
        
        return jsonify({
            'success': True,
            'message': f'Task {task_id} is now being processed',
            'task_status': 'processing'
        })
        
    except Exception as e:
        app.logger.error(f"Error processing task now: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/process-all', methods=['POST'])
@dual_auth_required
def process_all_tasks():
    """Process all pending tasks immediately"""
    try:
        # Get all pending tasks
        pending_tasks = [task for task in analysis_queue.tasks.values() if task.status == 'pending']
        
        if not pending_tasks:
            return jsonify({
                'success': False,
                'error': 'No pending tasks found'
            }), 404
        
        # Process each task
        for task in pending_tasks:
            task.priority = True  # Set all tasks to high priority
        
        # Trigger processing to start
        analysis_queue._enable_processing()
        
        return jsonify({
            'success': True,
            'message': f'Processing started for {len(pending_tasks)} tasks',
            'task_count': len(pending_tasks)
        })
        
    except Exception as e:
        app.logger.error(f"Error processing all tasks: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500  

@app.route('/api/task/<task_id>/cancel', methods=['POST'])
@dual_auth_required
def cancel_task(task_id):
    """Cancel a pending task"""
    try:
        # Get the task
        task = analysis_queue.get_task(task_id)
        if not task:
            return jsonify({'success': False, 'error': 'Task not found'}), 404
        
        # Check if task can be cancelled (only pending tasks can be cancelled)
        if task.status != 'pending':
            return jsonify({
                'success': False, 
                'error': f'Task cannot be cancelled. Current status: {task.status}'
            }), 400
        
        # Cancel the task
        success = analysis_queue.cancel_task(task_id)
        
        if success:
            # Update the metadata file explicitly
            try:
                metadata_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{task.filename}.json")
                if os.path.exists(metadata_path):
                    with open(metadata_path, 'r') as f:
                        metadata = json.load(f)
                    
                    # Update status in metadata
                    metadata['status'] = 'cancelled'
                    
                    with open(metadata_path, 'w') as f:
                        json.dump(metadata, f, indent=2)
                        
                    app.logger.info(f"Updated metadata file for cancelled task {task_id}")
            except Exception as e:
                app.logger.error(f"Error updating metadata file: {str(e)}")
            
            return jsonify({
                'success': True,
                'message': f'Task {task_id} has been cancelled',
                'task_status': 'cancelled'
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Failed to cancel task'
            }), 500
            
    except Exception as e:
        app.logger.error(f"Error cancelling task: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500      

# New endpoint to generate Word report
@app.route('/api/generate-report', methods=['POST'])
@dual_auth_required
def generate_report():
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create a new Document
        doc = Document()
        
        # Add a title
        title = doc.add_heading('Hospital Cleanliness Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date information
        doc.add_paragraph(f'Report generated on: {get_current_time()}')
        doc.add_paragraph(f'Generated by: {get_current_username()}')
        
        # Add content from the request
        if request.is_json:
            report_data = request.json
            hospital_name = report_data.get('hospital_name', 'All Hospitals')
            date_range = report_data.get('date_range', 'All Time')
            images = report_data.get('images', [])
            
            # Add report metadata
            doc.add_heading('Report Information', level=1)
            doc.add_paragraph(f'Hospital: {hospital_name}')
            doc.add_paragraph(f'Date Range: {date_range}')
            
            # Add summary statistics
            doc.add_heading('Summary Statistics', level=1)
            total_images = len(images)
            doc.add_paragraph(f'Total Images Analyzed: {total_images}')
            
            if total_images > 0:
                # Calculate average scores
                avg_cleanliness = sum(img.get('analysis_results', {}).get('cleanliness', {}).get('score', 0) for img in images if img.get('analysis_results')) / total_images
                avg_hygiene = sum(img.get('analysis_results', {}).get('hygiene', {}).get('score', 0) for img in images if img.get('analysis_results')) / total_images
                avg_pest_control = sum(img.get('analysis_results', {}).get('pest_control', {}).get('score', 0) for img in images if img.get('analysis_results')) / total_images
                avg_overall = sum(img.get('analysis_results', {}).get('overall', {}).get('score', 0) for img in images if img.get('analysis_results')) / total_images
                
                doc.add_paragraph(f'Average Cleanliness Score: {avg_cleanliness:.2f}%')
                doc.add_paragraph(f'Average Hygiene Score: {avg_hygiene:.2f}%')
                doc.add_paragraph(f'Average Pest Control Score: {avg_pest_control:.2f}%')
                doc.add_paragraph(f'Average Overall Score: {avg_overall:.2f}%')
            
            # Add detailed analysis for each image
            doc.add_heading('Detailed Analysis', level=1)
            for i, img in enumerate(images):
                if 'analysis_results' in img:
                    doc.add_heading(f'Image {i+1}: {img.get("hospital_name")} - {img.get("category")}', level=2)
                    doc.add_paragraph(f'Location: {img.get("building", "Unknown")} - {img.get("floor", "Unknown")} - {img.get("location", "Unknown")}')
                    doc.add_paragraph(f'Timestamp: {img.get("timestamp", "Unknown")}')
                    doc.add_paragraph(f'Priority: {"Yes" if img.get("priority", False) else "No"}')
                    
                    results = img.get('analysis_results', {})
                    if results:
                        for category in ['cleanliness', 'hygiene', 'pest_control', 'overall']:
                            if category in results:
                                cat_data = results[category]
                                doc.add_paragraph(f'{category.capitalize()}: {cat_data.get("score", "N/A")}% - {cat_data.get("status", "N/A")}')
                        
                        doc.add_paragraph(f'Comment: {img.get("comment", "No comment available")}')
                    
                    # Add a paragraph break
                    doc.add_paragraph()
        
        # Save the document
        report_filename = f'hospital_report_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        report_path = os.path.join(app.config['UPLOAD_FOLDER'], report_filename)
        doc.save(report_path)
        
        # Return the document URL
        report_url = f'/uploads/{report_filename}'
        return jsonify({
            'success': True,
            'message': 'Report generated successfully',
            'report_url': report_url,
            'report_filename': report_filename
        })
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'success': False, 'error': 'Not found'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'success': False, 'error': 'Internal server error'}), 500

@app.errorhandler(401)
def unauthorized_error(error):
    return jsonify({'success': False, 'error': 'Unauthorized access'}), 401

# Cleanup function
def cleanup():
    try:
        if schedule_manager and schedule_manager.scheduler.running:
            schedule_manager.stop()
        analysis_queue.stop()
    except Exception as e:
        print(f"Cleanup error: {str(e)}")

# Register cleanup function
atexit.register(cleanup)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)